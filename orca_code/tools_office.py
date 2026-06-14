
import json
import logging
import sys
import threading
import time
from pathlib import Path

from orca_code.config import OUTPUT_DIR, ensure_pkg
from orca_code.utils import _validate_write_path

"""orca_code.tools_office — Excel, Word, screenshot, OCR."""


def read_excel(path: str, sheet_name: str = None) -> str:
    try:
        import openpyxl
    except ImportError:
        if ensure_pkg("openpyxl"):
            import openpyxl
        else:
            return "错误: 缺少 openpyxl (pip install openpyxl)"
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            rows.append(" | ".join(cells))
        wb.close()
        if not rows:
            return "(空工作表)"
        header = f"工作表: {ws.title} | 共 {len(rows)} 行\n"
        content = header + "\n".join(rows[:500])
        if len(rows) > 500:
            content += f"\n... (共 {len(rows)} 行，已截断)"
        return content
    except Exception as e:
        return f"错误: {e}"
def write_excel(path: str, data: str, sheet_name: str = "Sheet1") -> str:
    try:
        import openpyxl
    except ImportError:
        if ensure_pkg("openpyxl"):
            import openpyxl
        else:
            return "错误: 缺少 openpyxl (pip install openpyxl)"
    try:
        parsed = json.loads(data)
    except json.JSONDecodeError as e:
        return f"错误: JSON 解析失败 - {e}"
    p, error = _validate_write_path(path)
    if error:
        return error
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name[:31]
        if isinstance(parsed, list) and len(parsed) > 0:
            if isinstance(parsed[0], dict):
                headers = list(parsed[0].keys())
                ws.append(headers)
                for item in parsed:
                    ws.append([item.get(h, "") for h in headers])
            else:
                for row in parsed:
                    ws.append(list(row))
        wb.save(str(p))
        wb.close()
        return f"已写入 {path} (工作表: {sheet_name})"
    except Exception as e:
        logging.error(f"write_excel error: {e}")
        return f"错误: {e}"
def read_word(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        if ensure_pkg("python-docx", "docx"):
            from docx import Document
        else:
            return "错误: 缺少 python-docx (pip install python-docx)"
    try:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text:
            return "(空文档)"
        return text[:10000] if len(text) > 10000 else text
    except Exception as e:
        return f"错误: {e}"
def write_word(path: str, content: str, title: str = None) -> str:
    try:
        from docx import Document
    except ImportError:
        if ensure_pkg("python-docx", "docx"):
            from docx import Document
        else:
            return "错误: 缺少 python-docx (pip install python-docx)"
    p, error = _validate_write_path(path)
    if error:
        return error
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(p))
        return f"已写入 {path} ({len(content)} 字符)"
    except Exception as e:
        logging.error(f"write_word error: {e}")
        return f"错误: {e}"
def take_screenshot(window_title: str = None, save_path: str = None) -> str:
    try:
        from PIL import Image, ImageGrab
    except ImportError:
        if ensure_pkg("Pillow", "PIL"):
            from PIL import Image
        else:
            return "错误: 缺少 Pillow (pip install Pillow)"

    # Force all screenshots into output/ — don't let model scatter files across Desktop
    if save_path:
        p = Path(save_path)
        save_to = OUTPUT_DIR / p.name  # always just the filename, ignore any dir prefix
    else:
        save_to = OUTPUT_DIR / f"screenshot_{int(time.time())}.png"
    save_to.parent.mkdir(parents=True, exist_ok=True)

    if window_title and sys.platform == "win32":
        try:
            import win32con
            import win32gui
            import win32ui

            found_hwnd = None
            title_lower = window_title.lower()

            def _enum_cb(hwnd, _):
                nonlocal found_hwnd
                if found_hwnd:
                    return
                if win32gui.IsWindowVisible(hwnd):
                    t = win32gui.GetWindowText(hwnd)
                    if title_lower in t.lower():
                        found_hwnd = hwnd

            win32gui.EnumWindows(_enum_cb, None)
            if not found_hwnd:
                return f"错误: 未找到包含 '{window_title}' 的窗口"

            if win32gui.IsIconic(found_hwnd):
                win32gui.ShowWindow(found_hwnd, win32con.SW_RESTORE)
                time.sleep(0.3)

            left, top, right, bottom = win32gui.GetClientRect(found_hwnd)
            (left, top), (right, bottom) = (
                win32gui.ClientToScreen(found_hwnd, (left, top)),
                win32gui.ClientToScreen(found_hwnd, (right, bottom)),
            )
            width, height = right - left, bottom - top

            hwndDC = None
            mfcDC = None
            saveDC = None
            saveBitMap = None
            try:
                hwndDC = win32gui.GetWindowDC(found_hwnd)
                mfcDC = win32ui.CreateDCFromHandle(hwndDC)
                saveDC = mfcDC.CreateCompatibleDC()
                saveBitMap = win32ui.CreateBitmap()
                saveBitMap.CreateCompatibleBitmap(mfcDC, width, height)
                saveDC.SelectObject(saveBitMap)
                saveDC.BitBlt((0, 0), (width, height), mfcDC, (0, 0), win32con.SRCCOPY)

                bmpinfo = saveBitMap.GetInfo()
                bmpstr = saveBitMap.GetBitmapBits(True)
                im = Image.frombuffer(
                    "RGB", (bmpinfo["bmWidth"], bmpinfo["bmHeight"]),
                    bmpstr, "raw", "BGRX", 0, 1
                )
            finally:
                if saveBitMap:
                    win32gui.DeleteObject(saveBitMap.GetHandle())
                if saveDC:
                    saveDC.DeleteDC()
                if mfcDC:
                    mfcDC.DeleteDC()
                if hwndDC:
                    win32gui.ReleaseDC(found_hwnd, hwndDC)

            im.save(str(save_to))
            win_text = win32gui.GetWindowText(found_hwnd)
            return f"窗口截图已保存: {save_to} (窗口: {win_text}, {width}x{height})"
        except ImportError:
            return "错误: 窗口截图需要 pywin32 (pip install pywin32)"
        except Exception as e:
            return f"错误: 窗口截图失败 - {e}，尝试全屏截图..."

    # [C10] Use mss for faster fullscreen screenshots
    try:
        import mss
        save_to.parent.mkdir(parents=True, exist_ok=True)
        with mss.mss() as sct:
            monitor = sct.monitors[1]
            img = sct.grab(monitor)
            mss.tools.to_png(img.rgb, img.size, output=str(save_to))
        return f"全屏截图已保存: {save_to} ({monitor['width']}x{monitor['height']})"
    except ImportError:
        if ensure_pkg("mss"):
            import mss
            save_to.parent.mkdir(parents=True, exist_ok=True)
            with mss.mss() as sct:
                monitor = sct.monitors[1]
                img = sct.grab(monitor)
                mss.tools.to_png(img.rgb, img.size, output=str(save_to))
            return f"全屏截图已保存: {save_to} ({monitor['width']}x{monitor['height']})"
        return "错误: 缺少 mss (pip install mss)"
    except Exception as e:
        logging.error(f"截图失败: {e}")
        return f"错误: 截图失败 - {e}"
_ocr_lock = threading.Lock()
_ocr_engine = None
def ocr_image(path: str) -> str:
    global _ocr_engine
    try:
        if _ocr_engine is None:
            with _ocr_lock:
                if _ocr_engine is None:
                    try:
                        from rapidocr_onnxruntime import RapidOCR
                    except ImportError:
                        if ensure_pkg("rapidocr-onnxruntime", "rapidocr_onnxruntime"):
                            from rapidocr_onnxruntime import RapidOCR
                        else:
                            return "错误: 缺少 rapidocr-onnxruntime"
                    _ocr_engine = RapidOCR()
        p = Path(path)
        if not p.exists():
            return f"错误: 文件不存在 - {path}"
        result, elapse = _ocr_engine(str(p))
        if not result:
            return "未识别到文字"
        lines = [f"{i}. {text} (置信度: {score:.2f})" for i, (_, text, score) in enumerate(result, 1)]
        # elapse may be float or list[float]
        if isinstance(elapse, (list, tuple)):
            elapse_val = sum(elapse)
        else:
            elapse_val = float(elapse)
        lines.append(f"\n耗时: {elapse_val:.2f}s")
        return "\n".join(lines)
    except Exception as e:
        logging.error(f"OCR error: {e}")
        return f"错误: OCR 识别失败 - {e}"
